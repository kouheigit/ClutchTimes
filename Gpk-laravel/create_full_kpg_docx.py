#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Kpg実装手順書（全フェーズ詳細版）をdocxファイルとして生成
コードに色付け機能付き、全Phase 0-16を含む
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_code_block_colored(doc, code, language='php', description_before=None, description_after=None):
    """コードブロックを色付けして追加（改善版・詳細説明付き）"""
    
    # コード前の説明を追加
    if description_before:
        p_desc = doc.add_paragraph(description_before)
        p_desc.style = 'Normal'
        p_desc.runs[0].font.size = Pt(10)
        p_desc.runs[0].font.color.rgb = RGBColor(50, 50, 50)
        # 説明の後の余白を最小限に（0に設定）
        p_desc.paragraph_format.space_after = Pt(0)
    
    # コードブロックを枠で囲む（視覚的に見やすく）
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.2)
    # 見出しや説明文の直後のコードブロックは余白を0に
    p.paragraph_format.space_before = Pt(0)  # 余白を0に設定
    p.paragraph_format.space_after = Pt(0)  # コードブロック後の余白も0に設定（空行を完全に削除）
    p.paragraph_format.keep_together = True  # コードブロックを分割しない
    
    if language == 'php':
        php_keywords = [
            '<?php', 'namespace', 'use', 'class', 'function', 'public', 'private', 
            'protected', 'return', 'if', 'else', 'elseif', 'foreach', 'for', 'while', 
            'try', 'catch', 'throw', 'new', 'extends', 'implements', 'const', 'static',
            'abstract', 'final', 'interface', 'trait', 'array', 'string', 'int', 'bool',
            'null', 'true', 'false', 'this', 'self', 'parent', 'DB', 'Auth', 'Carbon',
            'use', 'Illuminate', 'App', 'namespace'
        ]
        
        # コードを整形（インデントを保持、空行を最適化）
        lines = code.split('\n')
        
        # 先頭の連続する空行を削除
        while lines and not lines[0].strip():
            lines.pop(0)
        # 末尾の連続する空行を削除
        while lines and not lines[-1].strip():
            lines.pop()
        
        # 連続する空行を1つに統一（ただし、クラス定義や関数定義の前後の空行は保持）
        optimized_lines = []
        prev_empty = False
        prev_line_stripped = ''
        
        for i, line in enumerate(lines):
            is_empty = not line.strip()
            current_stripped = line.strip()
            
            if is_empty:
                # 空行の場合
                if not prev_empty:
                    # 前の行が空行でない場合のみ追加（連続空行を1つに）
                    optimized_lines.append('')
                prev_empty = True
            else:
                # コード行の場合
                # クラス定義や関数定義の前の空行は保持
                is_class_or_function = (
                    current_stripped.startswith('class ') or
                    current_stripped.startswith('function ') or
                    current_stripped.startswith('public function ') or
                    current_stripped.startswith('private function ') or
                    current_stripped.startswith('protected function ') or
                    current_stripped.startswith('namespace ') or
                    current_stripped.startswith('use ') and current_stripped.endswith(';')
                )
                
                # 前の行がクラス/関数定義の終わりで、現在の行がクラス/関数定義の始まりの場合、空行を保持
                if (prev_line_stripped.endswith('}') or prev_line_stripped.endswith(';')) and is_class_or_function:
                    if prev_empty and optimized_lines and optimized_lines[-1] == '':
                        # 既に空行がある場合はそのまま
                        pass
                    elif not prev_empty:
                        # 前の行が空行でない場合、空行を追加
                        optimized_lines.append('')
                
                optimized_lines.append(line)
                prev_empty = False
                prev_line_stripped = current_stripped
        
        # 末尾の連続する空行を削除
        while optimized_lines and not optimized_lines[-1].strip():
            optimized_lines.pop()
        
        lines = optimized_lines
        
        for line_num, line in enumerate(lines):
            # 空行はそのまま追加
            if not line.strip():
                run = p.add_run('\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                continue
            
            # コメント行（// または # で始まる）
            stripped = line.strip()
            if stripped.startswith('//') or stripped.startswith('#'):
                run = p.add_run(line + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)  # グレー
                continue
            
            # 行をトークンに分割して色付け（改善版）
            i = 0
            while i < len(line):
                # 文字列リテラル（"..." または '...'）
                if line[i] == '"' or line[i] == "'":
                    quote_char = line[i]
                    start = i
                    i += 1
                    escaped = False
                    while i < len(line):
                        if escaped:
                            escaped = False
                            i += 1
                            continue
                        if line[i] == '\\':
                            escaped = True
                            i += 1
                            continue
                        if line[i] == quote_char:
                            i += 1
                            break
                        i += 1
                    
                    run = p.add_run(line[start:i])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 128, 0)  # 緑（文字列）
                    continue
                
                # キーワード検出（より正確に）
                matched = False
                for keyword in php_keywords:
                    keyword_len = len(keyword)
                    if (i + keyword_len <= len(line) and 
                        line[i:i+keyword_len] == keyword and
                        (i + keyword_len >= len(line) or 
                         not (line[i+keyword_len].isalnum() or line[i+keyword_len] == '_'))):
                        run = p.add_run(keyword)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 255)  # 青（キーワード）
                        i += keyword_len
                        matched = True
                        break
                
                if not matched:
                    # 通常の文字
                    run = p.add_run(line[i])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黒
                    i += 1
            
            run = p.add_run('\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # コード後の説明を追加
        if description_after:
            p_after = doc.add_paragraph()
            p_after.add_run('【補足説明】').bold = True
            p_after.add_run(' ' + description_after)
            p_after.runs[0].font.size = Pt(9)
            p_after.runs[1].font.size = Pt(9)
            p_after.runs[1].font.color.rgb = RGBColor(80, 80, 80)
            p_after.paragraph_format.space_after = Pt(0)  # 説明文の後の余白を0に
            
    elif language == 'bash':
        # Bashコマンドを整形して追加
        lines = code.split('\n')
        for line in lines:
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            if line.strip().startswith('#'):
                run.font.color.rgb = RGBColor(128, 128, 128)  # コメントはグレー
            else:
                run.font.color.rgb = RGBColor(0, 128, 0)  # コマンドは緑
        if description_after:
            p_after = doc.add_paragraph()
            p_after.add_run('【実行結果】').bold = True
            p_after.add_run(' ' + description_after)
            p_after.runs[0].font.size = Pt(9)
            p_after.runs[1].font.size = Pt(9)
            p_after.runs[1].font.color.rgb = RGBColor(80, 80, 80)
            
    elif language == 'yaml' or language == 'yml':
        # YAMLを整形
        lines = code.split('\n')
        indent_level = 0
        for line in lines:
            stripped = line.lstrip()
            if stripped.startswith('#'):
                run = p.add_run(line + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)  # コメントはグレー
            else:
                run = p.add_run(line + '\n')
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 255)  # YAMLは青
        if description_after:
            p_after = doc.add_paragraph()
            p_after.add_run('【設定説明】').bold = True
            p_after.add_run(' ' + description_after)
            p_after.runs[0].font.size = Pt(9)
            p_after.runs[1].font.size = Pt(9)
            p_after.runs[1].font.color.rgb = RGBColor(80, 80, 80)
            
    elif language == 'blade':
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 0, 128)  # 紫
    else:
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黒
    
    return p

def parse_markdown_to_docx(doc, markdown_content):
    """Markdownコンテンツをdocxに変換（詳細説明付き・空行最小化）"""
    lines = markdown_content.split('\n')
    i = 0
    prev_was_heading = False  # 前の行が見出しかどうか
    
    while i < len(lines):
        line = lines[i]
        
        # 見出し（GLAMDAY STYLEフォント適用）
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], 1)
            # GLAMDAY STYLEフォントを適用
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            # 次の行がコードブロックまたは空行+コードブロックの場合は余白を0に
            next_is_code = (i + 1 < len(lines) and lines[i + 1].startswith('```')) or \
                          (i + 2 < len(lines) and not lines[i + 1].strip() and lines[i + 2].startswith('```'))
            heading.paragraph_format.space_after = Pt(0) if next_is_code else Pt(2)
            prev_was_heading = True
        elif line.startswith('## '):
            heading_text = line[3:]
            heading = doc.add_heading(heading_text, 2)
            # GLAMDAY STYLEフォントを適用
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            # 次の行がコードブロックまたは空行+コードブロックの場合は余白を0に
            next_is_code = (i + 1 < len(lines) and lines[i + 1].startswith('```')) or \
                          (i + 2 < len(lines) and not lines[i + 1].strip() and lines[i + 2].startswith('```'))
            heading.paragraph_format.space_after = Pt(0) if next_is_code else Pt(1)
            prev_was_heading = True
            # 見出しの後に説明を追加（可能な場合）
            if i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].startswith('#') and not lines[i + 1].startswith('```'):
                next_line = lines[i + 1].strip()
                if not next_line.startswith('-') and not next_line.startswith('*'):
                    p = doc.add_paragraph(next_line)
                    p.runs[0].font.size = Pt(10)
                    p.runs[0].font.color.rgb = RGBColor(60, 60, 60)
                    # 次の行がコードブロックの場合は余白を0に
                    if i + 2 < len(lines) and lines[i + 2].startswith('```'):
                        p.paragraph_format.space_after = Pt(0)
                    else:
                        p.paragraph_format.space_after = Pt(1)
                    i += 1
                    prev_was_heading = False
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], 3)
            # GLAMDAY STYLEフォントを適用
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            # 次の行がコードブロックまたは空行+コードブロックの場合は余白を0に
            next_is_code = (i + 1 < len(lines) and lines[i + 1].startswith('```')) or \
                          (i + 2 < len(lines) and not lines[i + 1].strip() and lines[i + 2].startswith('```'))
            heading.paragraph_format.space_after = Pt(0) if next_is_code else Pt(1)
            prev_was_heading = True
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], 4)
            # GLAMDAY STYLEフォントを適用
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            # 次の行がコードブロックまたは空行+コードブロックの場合は余白を0に
            next_is_code = (i + 1 < len(lines) and lines[i + 1].startswith('```')) or \
                          (i + 2 < len(lines) and not lines[i + 1].strip() and lines[i + 2].startswith('```'))
            heading.paragraph_format.space_after = Pt(0) if next_is_code else Pt(0)
            prev_was_heading = True
        
        # コードブロック開始
        elif line.startswith('```'):
            lang = line[3:].strip() if len(line) > 3 else ''
            code_lines = []
            start_line_idx = i
            i += 1
            
            # コードブロックの前の説明を取得（直前の行のみ確認、空行はスキップ）
            description_before = None
            # 直前の行から遡って最初の非空行を取得
            for j in range(start_line_idx - 1, max(-1, start_line_idx - 4), -1):
                if j < 0:
                    break
                prev_line = lines[j].strip()
                if prev_line and not prev_line.startswith('#') and not prev_line.startswith('-') and not prev_line.startswith('*') and not prev_line.startswith('```'):
                    description_before = prev_line
                    break  # 最初に見つかった説明文を使用
            
            # 直前の段落の余白を0に設定（コードブロック直前の空行を完全に削除）
            # これはdocxの段落スタイルで制御するため、ここでは処理しない
            
            # コードブロックの内容を取得（終了記号を見つけるまで）
            found_end = False
            while i < len(lines):
                current_line = lines[i]
                # 終了記号をチェック（行全体が```で始まる場合のみ）
                if current_line.strip() == '```' or (current_line.startswith('```') and len(current_line.strip()) <= 3):
                    found_end = True
                    break
                # コード行を追加
                code_lines.append(current_line)
                i += 1
            
            # 終了記号が見つからない場合の警告
            if not found_end:
                print(f"⚠️  警告: コードブロックが正しく閉じられていません（行 {start_line_idx + 1}）")
            
            code = '\n'.join(code_lines)
            
            # コードブロックの後の説明を取得（次の数行を確認）
            description_after = None
            if found_end and i + 1 < len(lines):
                # 次の3行まで確認
                for j in range(i + 1, min(i + 4, len(lines))):
                    next_line = lines[j].strip()
                    if next_line and not next_line.startswith('#') and not next_line.startswith('```') and not next_line.startswith('-') and not next_line.startswith('*'):
                        if not description_after:
                            description_after = next_line
                        elif len(next_line) > len(description_after):
                            description_after = next_line
            
            # コードが空でない場合のみ追加
            if code.strip():
                # 長いコードブロックの場合は警告を出す
                if len(code_lines) > 100:
                    print(f"📝 長いコードブロックを処理中: {len(code_lines)}行（行 {start_line_idx + 1}-{i + 1}）")
                
                # 直前の段落の余白を0に設定するため、最後の段落を取得
                # 直前がコードブロックの場合は余白を0に
                if len(doc.paragraphs) > 0:
                    last_para = doc.paragraphs[-1]
                    # 直前の段落が空行の場合、削除（docxでは直接削除できないため、余白を0に）
                    if not last_para.text.strip():
                        last_para.paragraph_format.space_after = Pt(0)
                        last_para.paragraph_format.space_before = Pt(0)
                    # 直前の段落がコードブロックまたは見出しの場合、余白を0に
                    elif last_para.text.strip() or start_line_idx > 0:
                        # 直前の行がコードブロック終了記号（```）か確認
                        prev_is_code_end = start_line_idx > 0 and lines[start_line_idx - 1].strip() == '```'
                        # 直前の行が見出しか確認
                        prev_is_heading = start_line_idx > 0 and lines[start_line_idx - 1].startswith('#')
                        if prev_is_code_end or prev_is_heading:
                            last_para.paragraph_format.space_after = Pt(0)
                    else:
                        last_para.paragraph_format.space_after = Pt(0)
                
                add_code_block_colored(doc, code, lang, description_before, description_after)
            else:
                print(f"⚠️  警告: 空のコードブロックをスキップしました（行 {start_line_idx + 1}）")
            
            # 終了記号の行もスキップ（次のループで処理される）
            if found_end:
                # iは既に```の行を指しているので、そのまま次へ進む
                # コードブロック直後の空行をスキップ
                while i + 1 < len(lines) and not lines[i + 1].strip():
                    i += 1
        
        # リスト
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line[2:]
            # リスト項目に詳細説明を追加
            if '：' in list_text or ':' in list_text:
                parts = list_text.split('：' if '：' in list_text else ':', 1)
                p = doc.add_paragraph()
                p.add_run(parts[0]).bold = True
                if len(parts) > 1:
                    p.add_run('：' + parts[1])
                p.style = 'List Bullet'
                p.paragraph_format.space_after = Pt(0)  # リスト項目の余白を0に
            else:
                p = doc.add_paragraph(list_text, style='List Bullet')
                p.paragraph_format.space_after = Pt(0)  # リスト項目の余白を0に
            prev_was_heading = False
        
        # 番号付きリスト
        elif re.match(r'^\d+\.\s', line):
            list_text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(list_text, style='List Number')
            p.paragraph_format.space_after = Pt(0)  # リスト項目の余白を0に
            prev_was_heading = False
        
        # 空行（すべての空行を完全にスキップ）
        elif not line.strip():
            # 前の行を確認
            prev_line_empty = i > 0 and not lines[i-1].strip()
            # 次の行を確認
            next_line_code = i + 1 < len(lines) and lines[i + 1].startswith('```')
            # 前の行がコードブロック終了記号か確認
            prev_line_code_end = i > 0 and lines[i-1].strip() == '```'
            # 前の行が見出しの場合
            prev_line_heading = prev_was_heading or (i > 0 and (lines[i-1].startswith('#') or lines[i-1].startswith('####') or lines[i-1].startswith('###') or lines[i-1].startswith('##')))
            
            # コードブロックの直前の空行は常にスキップ
            if next_line_code:
                pass  # 空行を完全にスキップ
            # コードブロックの直後の空行もスキップ
            elif prev_line_code_end:
                pass  # 空行を完全にスキップ
            # 見出しの直後の空行もスキップ（見出しとコードブロックの間）
            elif prev_line_heading and next_line_code:
                pass  # 空行を完全にスキップ
            # 見出しの直後の空行はスキップ
            elif prev_line_heading:
                pass  # 見出し直後の空行をスキップ
            # 連続する空行は追加しない
            elif prev_line_empty:
                pass  # 連続空行をスキップ
            # それ以外の場合は空行を追加しない（空行を完全に削除）
            else:
                pass  # 空行を完全にスキップ
            
            prev_was_heading = False
        
        # 通常のテキスト
        elif line.strip() and not line.startswith('#'):
            # 重要な説明文を太字にする
            if line.strip().startswith('⚠️') or line.strip().startswith('注意') or line.strip().startswith('重要'):
                p = doc.add_paragraph()
                p.add_run(line.strip()).bold = True
                p.runs[0].font.color.rgb = RGBColor(200, 0, 0)  # 赤
                p.paragraph_format.space_after = Pt(0)  # 余白を0に
            else:
                p = doc.add_paragraph(line)
                # 次の行がコードブロックの場合は余白を0に
                if i + 1 < len(lines) and lines[i + 1].startswith('```'):
                    p.paragraph_format.space_after = Pt(0)
                else:
                    p.paragraph_format.space_after = Pt(0)  # 通常のテキストも余白を0に
            prev_was_heading = False
        
        i += 1

def setup_document_styles(doc):
    """GLAMDAY STYLEデザインに基づいたドキュメントスタイルを設定"""
    # スタイルをカスタマイズ
    styles = doc.styles
    
    # 見出し1のスタイル（タイトル用）
    if 'Heading 1' in styles:
        h1_style = styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'  # セリフフォント
        h1_font.size = Pt(28)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)  # 黒
    
    # 見出し2のスタイル
    if 'Heading 2' in styles:
        h2_style = styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(18)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
    
    # 見出し3のスタイル
    if 'Heading 3' in styles:
        h3_style = styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 0, 0)
    
    # 通常テキストのスタイル
    if 'Normal' in styles:
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Hiragino Sans'  # 日本語用フォント
        normal_font.size = Pt(10.5)
        normal_font.color.rgb = RGBColor(0, 0, 0)

def create_full_documentation():
    """全フェーズを含む完全な実装手順書を作成（GLAMDAY STYLEデザイン適用）"""
    
    doc = Document()
    
    # ドキュメントスタイルを設定
    setup_document_styles(doc)
    
    # タイトルページ - GLAMDAY STYLEデザイン
    # GSロゴ風のテキスト（モノグラム風）
    gs_logo = doc.add_paragraph()
    gs_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gs_run = gs_logo.add_run('GS')
    gs_run.font.name = 'Times New Roman'
    gs_run.font.size = Pt(48)
    gs_run.font.bold = True
    gs_run.font.color.rgb = RGBColor(0, 0, 0)
    gs_logo.paragraph_format.space_after = Pt(8)
    
    # GLAMDAY STYLE タイトル
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('GLAMDAY STYLE')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(20)
    
    # OFFICIAL OWNER'S SITE
    official = doc.add_paragraph()
    official.alignment = WD_ALIGN_PARAGRAPH.CENTER
    official_run = official.add_run('OFFICIAL')
    official_run.font.name = 'Times New Roman'
    official_run.font.size = Pt(14)
    official_run.font.italic = True
    official_run.font.color.rgb = RGBColor(0, 0, 0)
    official.paragraph_format.space_after = Pt(4)
    
    owners_site = doc.add_paragraph()
    owners_site.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owners_run = owners_site.add_run("OWNER'S SITE")
    owners_run.font.name = 'Times New Roman'
    owners_run.font.size = Pt(14)
    owners_run.font.italic = True
    owners_run.font.color.rgb = RGBColor(0, 0, 0)
    owners_site.paragraph_format.space_after = Pt(30)
    
    # KGP実装手順書
    kpg_title = doc.add_paragraph()
    kpg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kpg_run = kpg_title.add_run('KGP実装手順書')
    kpg_run.font.name = 'Hiragino Sans'
    kpg_run.font.size = Pt(20)
    kpg_run.font.bold = True
    kpg_run.font.color.rgb = RGBColor(0, 0, 0)
    kpg_title.paragraph_format.space_after = Pt(8)
    
    # サブタイトル
    subtitle = doc.add_paragraph('会員制宿泊予約システム 完全実装ガイド（全フェーズ詳細版）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.name = 'Hiragino Sans'
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(60, 60, 60)
    
    # 右下に日本語テキスト
    doc.add_paragraph()  # スペース
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_text.add_run('オーナー様専用サイト\n総合ページ')
    footer_run.font.name = 'Hiragino Sans'
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()
    
    # 目次（GLAMDAY STYLEフォント適用）
    toc_heading = doc.add_heading('📚 目次', 1)
    for run in toc_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    toc_items = [
        'システム概要',
        'Phase 0: プロジェクト準備（3-5日）',
        'Phase 1: データベース設計（1週間）',
        'Phase 2: モデル・定数定義（3-4日）',
        'Phase 3: 認証・基本機能（1週間）',
        'Phase 4: 予約システムコア実装（2週間）',
        'Phase 5: サービス注文・カート機能（1週間）',
        'Phase 6: 決済連携（Veritrans）（1-2週間）',
        'Phase 7: ポイントシステム（1週間）',
        'Phase 8: 招待機能（3-5日）',
        'Phase 9: 管理画面実装（2週間）',
        'Phase 10: フロント画面実装（2-3週間）',
        'Phase 11: バリデーション・Request（1週間）',
        'Phase 12: メール機能（3-5日）',
        'Phase 13: 外部API連携（3-5日）',
        'Phase 14: SPA化対応（2-3週間）',
        'Phase 15: テスト・品質向上（2週間）',
        'Phase 16: デプロイ・運用（1週間）',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Phase 0の内容を手動で追加（詳細説明付き・GLAMDAY STYLEフォント適用）
    phase0_heading = doc.add_heading('Phase 0: プロジェクト準備（3-5日）', 1)
    for run in phase0_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    goal_heading = doc.add_heading('目標', 2)
    for run in goal_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p_goal = doc.add_paragraph()
    p_goal.add_run('このフェーズでは、開発環境の基盤を構築します。具体的には以下の作業を行います：').bold = True
    doc.add_paragraph('• Docker環境の構築（PHP、Nginx、MySQL、phpMyAdmin、MailHog）', style='List Bullet')
    doc.add_paragraph('• Laravelプロジェクトの初期設定', style='List Bullet')
    doc.add_paragraph('• 必要パッケージのインストール（Laravel Breeze、Laravel Admin、Veritrans SDKなど）', style='List Bullet')
    doc.add_paragraph('• 開発環境の動作確認', style='List Bullet')
    
    doc.add_paragraph('このフェーズを完了することで、以降の開発作業を効率的に進めることができます。')
    
    step01_heading = doc.add_heading('Step 0-1: Docker環境構築', 2)
    for run in step01_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('Docker Composeを使用して、複数のコンテナを一元管理できる環境を構築します。')
    doc.add_paragraph('Docker Composeの利点：')
    doc.add_paragraph('• 複数のコンテナを1つのコマンドで起動・停止できる', style='List Bullet')
    doc.add_paragraph('• 環境変数で設定を一元管理できる', style='List Bullet')
    doc.add_paragraph('• チームメンバー間で同じ環境を共有できる', style='List Bullet')
    doc.add_paragraph('• 本番環境に近い構成で開発できる', style='List Bullet')
    
    dir_heading = doc.add_heading('1. ディレクトリ構成作成', 3)
    for run in dir_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('プロジェクトのディレクトリ構成を作成します。この構成により、設定ファイルとソースコードを適切に分離できます。')
    add_code_block_colored(
        doc, 
        '''mkdir -p kpg-laravel/infra/docker/nginx
mkdir -p kpg-laravel/infra/docker/php
mkdir -p kpg-laravel/infra/docker/mysql
mkdir -p kpg-laravel/src''', 
        'bash',
        description_before='以下のコマンドを実行して、プロジェクトのディレクトリ構造を作成します。',
        description_after='このコマンドにより、infraディレクトリ配下にDocker設定ファイル用のディレクトリ、srcディレクトリにLaravelアプリケーション本体が配置されます。'
    )
    
    doc.add_paragraph('作成されるディレクトリ構造の説明：')
    doc.add_paragraph('• infra/docker/nginx: Nginx設定ファイル（Webサーバー設定、リバースプロキシ設定など）', style='List Bullet')
    doc.add_paragraph('• infra/docker/php: PHP設定ファイル（PHP-FPM設定、拡張機能設定など）', style='List Bullet')
    doc.add_paragraph('• infra/docker/mysql: MySQL設定ファイル（データベース設定、文字コード設定など）', style='List Bullet')
    doc.add_paragraph('• src: Laravelアプリケーション本体（すべてのソースコードがここに配置されます）', style='List Bullet')
    
    docker_heading = doc.add_heading('2. docker-compose.yml 作成', 3)
    for run in docker_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph('プロジェクトルートにdocker-compose.ymlを作成します。このファイルは、複数のDockerコンテナを管理するための設定ファイルです。')
    doc.add_paragraph('docker-compose.ymlの主な役割：')
    doc.add_paragraph('• サービスの定義（app、web、dbなど）', style='List Bullet')
    doc.add_paragraph('• コンテナ間のネットワーク設定', style='List Bullet')
    doc.add_paragraph('• ボリュームマウント設定', style='List Bullet')
    doc.add_paragraph('• 環境変数の設定', style='List Bullet')
    
    add_code_block_colored(
        doc,
        '''version: "3.9"
volumes:
  db-store:
  psysh-store:
services:
  app:
    build:
      context: .
      dockerfile: ./infra/docker/php/Dockerfile
      target: ${APP_BUILD_TARGET:-development}
    volumes:
      - type: bind
        source: ./src
        target: /data
    environment:
      - APP_DEBUG=${APP_DEBUG:-true}
      - DB_CONNECTION=${DB_CONNECTION:-mysql}
      - DB_HOST=${DB_HOST:-db}
      - DB_PORT=${DB_PORT:-3306}
      - DB_DATABASE=${DB_DATABASE:-laravel}
      - DB_USERNAME=${DB_USERNAME:-phper}
      - DB_PASSWORD=${DB_PASSWORD:-secret}

  web:
    build:
      context: .
      dockerfile: ./infra/docker/nginx/Dockerfile
    ports:
      - "8081:80"
    volumes:
      - type: bind
        source: ./src
        target: /data

  db:
    build:
      context: .
      dockerfile: ./infra/docker/mysql/Dockerfile
    ports:
      - "3306:3306"
    volumes:
      - type: volume
        source: db-store
        target: /var/lib/mysql
    environment:
      - MYSQL_DATABASE=${DB_DATABASE:-laravel}
      - MYSQL_USER=${DB_USERNAME:-phper}
      - MYSQL_PASSWORD=${DB_PASSWORD:-secret}
      - MYSQL_ROOT_PASSWORD=${DB_PASSWORD:-secret}

  phpmyadmin:
    image: phpmyadmin:5.2
    environment:
      - PMA_HOST=db
    ports:
      - "8080:80"

  mailhog:
    image: mailhog/mailhog
    ports:
      - "8025:8025"''',
        'yaml',
        description_before='プロジェクトルートにdocker-compose.ymlを作成し、以下の内容を記述します。',
        description_after='この設定により、5つのサービス（app、web、db、phpmyadmin、mailhog）が定義されます。各サービスの詳細は後述します。'
    )
    
    doc.add_paragraph('各サービスの詳細説明：')
    doc.add_paragraph('• app: PHP-FPMコンテナ（Laravelアプリケーション実行環境）。/dataディレクトリにソースコードがマウントされます。', style='List Bullet')
    doc.add_paragraph('• web: Nginxコンテナ（Webサーバー）。ポート8081でアクセス可能です。appコンテナのPHP-FPMにリクエストを転送します。', style='List Bullet')
    doc.add_paragraph('• db: MySQLコンテナ（データベース）。ポート3306でアクセス可能です。データはdb-storeボリュームに永続化されます。', style='List Bullet')
    doc.add_paragraph('• phpmyadmin: phpMyAdminコンテナ（データベース管理ツール）。ポート8080でアクセス可能です。dbコンテナに接続してデータベースを管理できます。', style='List Bullet')
    doc.add_paragraph('• mailhog: MailHogコンテナ（メールテストツール）。ポート8025でアクセス可能です。開発環境でメール送信をテストできます。', style='List Bullet')
    
    doc.add_page_break()
    
    # Markdownファイルを読み込んで追加
    md_files = [
        ('Kpg復元/kpg-laravel実装手順書_目次.md', 'システム概要'),
        ('Kpg復元/kpg-laravel実装手順書_Phase1-3.md', 'Phase 1-3'),
        ('Kpg復元/kpg-laravel実装手順書_Phase4-6.md', 'Phase 4-6'),
        ('Kpg復元/kpg-laravel実装手順書_Phase7-9.md', 'Phase 7-9'),
        ('Kpg復元/kpg-laravel実装手順書_Phase10-12.md', 'Phase 10-12'),
        ('Kpg復元/kpg-laravel実装手順書_Phase13-16.md', 'Phase 13-16'),
    ]
    
    for md_file, section_name in md_files:
        if os.path.exists(md_file):
            print(f'📖 {section_name} を読み込み中...')
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
                parse_markdown_to_docx(doc, content)
                doc.add_page_break()
        else:
            print(f'⚠️  {md_file} が見つかりません')
    
    # 保存
    output_path = '/Users/user/Desktop/pgfile/Gpk-laravel/KGP実装手順書.docx'
    doc.save(output_path)
    print(f'\n✅ KGP実装手順書（全フェーズ詳細版）を {output_path} に保存しました。')
    print(f'   全{len(md_files) + 1}セクション（Phase 0含む）を含む完全版です。')
    print(f'   • 詳細な説明を追加')
    print(f'   • コードを整形・色付け')
    print(f'   • 各ステップに補足説明を追加')

if __name__ == '__main__':
    create_full_documentation()

